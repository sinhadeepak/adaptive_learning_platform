#!/usr/bin/env python3
"""
md_to_docx.py — convert Aurora design-system Markdown files to DOCX.

Pandoc is not installed on this host, so this script renders Markdown via
python-docx with the subset of features used by the Aurora docs:

  * ATX headings (#..######)
  * Paragraphs
  * Bullet (- / *) and numbered lists (with 1 level of nesting)
  * GFM pipe tables
  * Fenced code blocks (``` ... ```)
  * Inline code (`x`)
  * Bold (**x**), italic (*x*), and links ([text](url))
  * Horizontal rules (---)
  * Block quotes (>) — rendered as italic indented paragraphs

ASCII wireframes inside fenced code blocks render correctly because the
script uses a monospace font (Consolas / Courier New) for code.

Usage:
    python3 md_to_docx.py FILE.md [FILE.md ...]
    python3 md_to_docx.py --all     # converts every .md in design-system tree

Each input FILE.md produces FILE.docx next to it.
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
HR_RE = re.compile(r"^\s*(---+|\*\*\*+|___+)\s*$")
FENCE_RE = re.compile(r"^```(\w*)\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?(\s*:?-+:?\s*\|)+\s*:?-+:?\s*\|?\s*$")
BULLET_RE = re.compile(r"^(\s*)([-*])\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")

INLINE_PATTERNS = [
    ("code", re.compile(r"`([^`]+?)`")),
    ("bold", re.compile(r"\*\*([^*]+?)\*\*")),
    ("italic", re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")),
    ("link", re.compile(r"\[([^\]]+)\]\(([^)]+)\)")),
]


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_inline(paragraph, text: str) -> None:
    """
    Render markdown inline syntax into runs on `paragraph`.

    We do a sequential, non-overlapping match pass: find the earliest match
    among inline patterns, emit text before it, emit the styled run, repeat.
    """
    while text:
        best = None
        best_kind = None
        for kind, pat in INLINE_PATTERNS:
            m = pat.search(text)
            if m and (best is None or m.start() < best.start()):
                best = m
                best_kind = kind

        if best is None:
            paragraph.add_run(text)
            return

        if best.start() > 0:
            paragraph.add_run(text[: best.start()])

        if best_kind == "code":
            run = paragraph.add_run(best.group(1))
            run.font.name = "Consolas"
            rpr = run._element.rPr
            rFonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
            if rFonts is not None:
                rFonts.set(qn("w:ascii"), "Consolas")
                rFonts.set(qn("w:hAnsi"), "Consolas")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F1F5F9")
            run._element.rPr.append(shd)
        elif best_kind == "bold":
            run = paragraph.add_run(best.group(1))
            run.bold = True
        elif best_kind == "italic":
            run = paragraph.add_run(best.group(1))
            run.italic = True
        elif best_kind == "link":
            add_hyperlink(paragraph, url=best.group(2), text=best.group(1))

        text = text[best.end():]


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), "CBD5E1")


def setup_styles(doc: Document) -> None:
    """Configure default fonts, colors, heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading_sizes = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}
    for level, size in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Hyperlink style
    if "Hyperlink" not in [s.name for s in doc.styles]:
        hl = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        hl.font.color.rgb = RGBColor(0x5B, 0x5B, 0xD6)
        hl.font.underline = True


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a fenced code block as a monospace, shaded paragraph."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.25)
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)

    # Add shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F8FAFC")
    pPr.append(shd)

    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rFonts)
    run.font.size = Pt(9)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(10)
        # header fill
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F1F5F9")
        tcPr.append(shd)
        set_cell_border(cell, top=True, left=True, bottom=True, right=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(header):
                continue
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            render_inline(p, cell_text.strip())
            for run in p.runs:
                run.font.size = Pt(10)
            set_cell_border(cell, top=True, left=True, bottom=True, right=True)


def parse_table(lines: list[str], start: int) -> tuple[int, list[str], list[list[str]]] | None:
    """If lines[start] starts a markdown table, parse and return (next_idx, header, rows)."""
    if start + 1 >= len(lines):
        return None
    header_line = lines[start]
    sep_line = lines[start + 1]
    if "|" not in header_line or not TABLE_SEP_RE.match(sep_line):
        return None
    header = [c.strip() for c in header_line.strip().strip("|").split("|")]
    rows: list[list[str]] = []
    i = start + 2
    while i < len(lines):
        line = lines[i]
        if "|" not in line or not line.strip():
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return i, header, rows


def convert_file(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    setup_styles(doc)

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_fence = False
    fence_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if in_fence:
            if FENCE_RE.match(line):
                add_code_block(doc, fence_lines)
                fence_lines = []
                in_fence = False
            else:
                fence_lines.append(line)
            i += 1
            continue
        m = FENCE_RE.match(line)
        if m:
            in_fence = True
            fence_lines = []
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CBD5E1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            p = doc.add_heading(level=min(level, 6))
            for r in p.runs:
                r.text = ""
            render_inline(p, heading_text)
            i += 1
            continue

        # Table
        tbl = parse_table(lines, i)
        if tbl:
            i, header, rows = tbl
            add_table(doc, header, rows)
            continue

        # Blockquote (single-paragraph)
        m = BLOCKQUOTE_RE.match(line)
        if m:
            quote_lines = [m.group(1)]
            j = i + 1
            while j < len(lines):
                m2 = BLOCKQUOTE_RE.match(lines[j])
                if not m2:
                    break
                quote_lines.append(m2.group(1))
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(" ".join(quote_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            i = j
            continue

        # Bullet list
        m = BULLET_RE.match(line)
        if m:
            indent = len(m.group(1))
            content = m.group(3)
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            render_inline(p, content)
            i += 1
            continue

        # Numbered list
        m = NUMBERED_RE.match(line)
        if m:
            indent = len(m.group(1))
            content = m.group(3)
            style = "List Number" if indent < 2 else "List Number 2"
            p = doc.add_paragraph(style=style)
            render_inline(p, content)
            i += 1
            continue

        # Blank line: skip
        if not line.strip():
            i += 1
            continue

        # Paragraph (may span continued lines until blank/break)
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if (
                HEADING_RE.match(nxt)
                or HR_RE.match(nxt)
                or FENCE_RE.match(nxt)
                or BULLET_RE.match(nxt)
                or NUMBERED_RE.match(nxt)
                or BLOCKQUOTE_RE.match(nxt)
                or "|" in nxt and j + 1 < len(lines) and TABLE_SEP_RE.match(lines[j + 1])
            ):
                break
            para_lines.append(nxt)
            j += 1
        p = doc.add_paragraph()
        render_inline(p, " ".join(para_lines))
        i = j

    # Trailing fence safety
    if in_fence and fence_lines:
        add_code_block(doc, fence_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("files", nargs="*", help="Markdown files to convert.")
    parser.add_argument(
        "--all",
        action="store_true",
        help="Convert all .md files in docs/02-design/ tree + the Aurora ADRs.",
    )
    args = parser.parse_args(argv)

    repo_root = Path(__file__).resolve().parents[3]
    paths: list[Path] = []

    if args.all:
        design_dir = repo_root / "docs" / "02-design"
        paths.extend(p for p in design_dir.glob("design-system-v2-aurora.md"))
        paths.extend(sorted((design_dir / "redesign").glob("*.md")))
        adr_dir = repo_root / "docs" / "adr"
        for slug in ("0028-design-system-v2-aurora.md", "0029-component-primitives-package.md"):
            p = adr_dir / slug
            if p.exists():
                paths.append(p)
    else:
        for f in args.files:
            p = Path(f).resolve()
            if not p.exists():
                print(f"warn: missing {p}", file=sys.stderr)
                continue
            paths.append(p)

    if not paths:
        parser.print_help()
        return 1

    for md in paths:
        docx = md.with_suffix(".docx")
        print(f"convert: {md.relative_to(repo_root)} -> {docx.relative_to(repo_root)}")
        convert_file(md, docx)

    print(f"done: {len(paths)} file(s) converted")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
